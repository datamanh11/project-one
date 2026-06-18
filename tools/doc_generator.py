#!/usr/bin/env python3
"""
Document generator for resume and cover letter output.
Creates DOCX files from structured content and converts to PDF using Word on macOS.
"""

import sys
import os
import json
import subprocess
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1, font_size=14, bold=True, color=(31, 73, 125)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=font_size, bold=bold, color=color)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    # Add bottom border to section headings
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1F497D')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_contact_header(doc, contact_info):
    """Add name and contact info at top of resume."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = p.add_run(contact_info.get("name", ""))
    set_font(name_run, size=18, bold=True, color=(31, 73, 125))

    contact_line = doc.add_paragraph()
    contact_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = []
    for field in ["phone", "email", "linkedin", "location"]:
        if contact_info.get(field):
            parts.append(contact_info[field])
    contact_run = contact_line.add_run(" | ".join(parts))
    set_font(contact_run, size=10, color=(89, 89, 89))
    contact_line.paragraph_format.space_after = Pt(4)


def create_resume_docx(output_path, resume_data):
    """
    resume_data structure:
    {
      "contact": {"name": ..., "phone": ..., "email": ..., "linkedin": ..., "location": ...},
      "summary": "...",
      "experience": [
        {"title": "...", "company": "...", "dates": "...", "bullets": ["..."]}
      ],
      "skills": {"category": ["skill1", "skill2"]},  # or just list
      "education": [
        {"degree": "...", "school": "...", "dates": "...", "details": "..."}
      ],
      "certifications": ["..."],
      "sections": [{"heading": "...", "content": "..."}]  # additional sections
    }
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Default paragraph spacing
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    contact = resume_data.get("contact", {})
    add_contact_header(doc, contact)

    if resume_data.get("summary"):
        add_heading(doc, "PROFESSIONAL SUMMARY")
        p = doc.add_paragraph(resume_data["summary"])
        p.paragraph_format.space_after = Pt(4)

    if resume_data.get("experience"):
        add_heading(doc, "PROFESSIONAL EXPERIENCE")
        for job in resume_data["experience"]:
            # Job title line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            title_run = p.add_run(job.get("title", ""))
            set_font(title_run, bold=True, size=11)
            if job.get("company"):
                sep_run = p.add_run("  |  ")
                set_font(sep_run, size=11, color=(89, 89, 89))
                co_run = p.add_run(job["company"])
                set_font(co_run, size=11, italic=True, color=(31, 73, 125))
            if job.get("dates"):
                dates_run = p.add_run(f"  •  {job['dates']}")
                set_font(dates_run, size=10, color=(89, 89, 89))
            if job.get("location"):
                loc_run = p.add_run(f"  •  {job['location']}")
                set_font(loc_run, size=10, color=(89, 89, 89))

            for bullet in job.get("bullets", []):
                bp = doc.add_paragraph(style='List Bullet')
                bp.paragraph_format.left_indent = Inches(0.25)
                bp.paragraph_format.space_after = Pt(1)
                brun = bp.add_run(bullet)
                set_font(brun, size=10.5)

    if resume_data.get("skills"):
        add_heading(doc, "SKILLS")
        skills = resume_data["skills"]
        if isinstance(skills, dict):
            for category, skill_list in skills.items():
                p = doc.add_paragraph()
                cat_run = p.add_run(f"{category}: ")
                set_font(cat_run, bold=True, size=10.5)
                skills_run = p.add_run(", ".join(skill_list))
                set_font(skills_run, size=10.5)
                p.paragraph_format.space_after = Pt(2)
        elif isinstance(skills, list):
            p = doc.add_paragraph(", ".join(skills))
            p.paragraph_format.space_after = Pt(2)

    if resume_data.get("education"):
        add_heading(doc, "EDUCATION")
        for edu in resume_data["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            deg_run = p.add_run(edu.get("degree", ""))
            set_font(deg_run, bold=True, size=11)
            if edu.get("school"):
                school_run = p.add_run(f"  —  {edu['school']}")
                set_font(school_run, size=11, italic=True)
            if edu.get("dates"):
                date_run = p.add_run(f"  •  {edu['dates']}")
                set_font(date_run, size=10, color=(89, 89, 89))
            if edu.get("details"):
                dp = doc.add_paragraph(edu["details"])
                set_font(dp.runs[0] if dp.runs else dp.add_run(""), size=10)
                dp.paragraph_format.space_after = Pt(2)

    if resume_data.get("certifications"):
        add_heading(doc, "CERTIFICATIONS")
        for cert in resume_data["certifications"]:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(cert)
            set_font(r, size=10.5)

    for extra_section in resume_data.get("sections", []):
        add_heading(doc, extra_section.get("heading", "").upper())
        p = doc.add_paragraph(extra_section.get("content", ""))
        p.paragraph_format.space_after = Pt(4)

    doc.save(output_path)
    print(f"Saved resume DOCX: {output_path}")


def create_cover_letter_docx(output_path, cl_data):
    """
    cl_data structure:
    {
      "sender": {"name": ..., "address": ..., "phone": ..., "email": ...},
      "date": "...",
      "recipient": {"name": ..., "title": ..., "company": ..., "address": ...},
      "subject": "Re: ...",
      "salutation": "Dear ...",
      "body_paragraphs": ["paragraph 1 text", "paragraph 2 text", ...],
      "closing": "Sincerely,",
      "signature": "Full Name"
    }
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    sender = cl_data.get("sender", {})
    if sender.get("name"):
        p = doc.add_paragraph()
        r = p.add_run(sender["name"])
        set_font(r, size=14, bold=True, color=(31, 73, 125))
    if sender.get("address"):
        p = doc.add_paragraph(sender["address"])
        p.paragraph_format.space_after = Pt(0)
    contact_parts = [v for k, v in sender.items() if k in ("phone", "email") and v]
    if contact_parts:
        p = doc.add_paragraph(" | ".join(contact_parts))
        p.paragraph_format.space_after = Pt(12)

    if cl_data.get("date"):
        p = doc.add_paragraph(cl_data["date"])
        p.paragraph_format.space_after = Pt(12)

    recipient = cl_data.get("recipient", {})
    for field in ["name", "title", "company", "address"]:
        if recipient.get(field):
            p = doc.add_paragraph(recipient[field])
            p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()

    if cl_data.get("subject"):
        p = doc.add_paragraph()
        r = p.add_run(cl_data["subject"])
        set_font(r, bold=True, size=11)
        p.paragraph_format.space_after = Pt(6)

    if cl_data.get("salutation"):
        p = doc.add_paragraph(cl_data["salutation"])
        p.paragraph_format.space_after = Pt(6)

    for para_text in cl_data.get("body_paragraphs", []):
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.first_line_indent = Pt(0)

    if cl_data.get("closing"):
        p = doc.add_paragraph(cl_data["closing"])
        p.paragraph_format.space_after = Pt(48)

    if cl_data.get("signature"):
        p = doc.add_paragraph(cl_data["signature"])
        r = p.runs[0] if p.runs else p.add_run(cl_data["signature"])
        set_font(r, bold=True, size=11)

    doc.save(output_path)
    print(f"Saved cover letter DOCX: {output_path}")


def convert_to_pdf_via_word(docx_path, pdf_path):
    """Convert DOCX to PDF using Microsoft Word on macOS via AppleScript."""
    docx_abs = str(Path(docx_path).resolve())
    pdf_abs = str(Path(pdf_path).resolve())

    script = f'''
    tell application "Microsoft Word"
        set theDoc to open POSIX file "{docx_abs}"
        set theRange to create range of theDoc start 0 end 0
        save as theDoc file name "{pdf_abs}" file format format PDF
        close theDoc saving no
    end tell
    '''
    result = subprocess.run(['osascript', '-e', script], capture_output=True, text=True)
    if result.returncode != 0:
        print(f"Word PDF conversion error: {result.stderr}", file=sys.stderr)
        # Fallback: try with just the file format approach
        script2 = f'''
        tell application "Microsoft Word"
            set theDoc to open POSIX file "{docx_abs}"
            save as theDoc file name "{pdf_abs}" file format format PDF
            close theDoc saving no
        end tell
        '''
        result2 = subprocess.run(['osascript', '-e', script2], capture_output=True, text=True)
        if result2.returncode != 0:
            print(f"Fallback PDF conversion also failed: {result2.stderr}", file=sys.stderr)
            return False
    print(f"Saved PDF: {pdf_abs}")
    return True


def sanitize_filename(name):
    """Remove or replace characters not safe for filenames."""
    return re.sub(r'[^\w\s\-]', '', name).strip().replace(' ', '_')


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: doc_generator.py <json_input_file>")
        sys.exit(1)

    with open(sys.argv[1]) as f:
        config = json.load(f)

    doc_type = config.get("type")  # "resume" or "cover_letter"
    output_docx = config.get("output_docx")
    output_pdf = config.get("output_pdf")
    data = config.get("data")

    if doc_type == "resume":
        create_resume_docx(output_docx, data)
    elif doc_type == "cover_letter":
        create_cover_letter_docx(output_docx, data)
    else:
        print(f"Unknown type: {doc_type}", file=sys.stderr)
        sys.exit(1)

    if output_pdf:
        convert_to_pdf_via_word(output_docx, output_pdf)
